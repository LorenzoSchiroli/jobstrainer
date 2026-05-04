"""Generate a tailored cover letter, CV, and LinkedIn messages for a job application."""
import json
import os
import re
import shutil
from docx import Document
import pyperclip
from groq import Groq
from dotenv import load_dotenv

load_dotenv()

CV_PATH = "data/lorenzo_schiroli_cv.docx"


def extract_cv_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def read_job_description() -> str:
    text = pyperclip.paste()
    if not text.strip():
        raise ValueError("Clipboard is empty. Copy the job description first.")
    return text.strip()


def get_client() -> Groq:
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        raise EnvironmentError(
            "GROQ_API_KEY environment variable not set. "
            "Get a free key at https://console.groq.com/keys"
        )
    return Groq(api_key=api_key)


def generate_cover_letter(cv_text: str, job_description: str, client: Groq) -> tuple[str, str, str]:
    prompt = (
        "You are an expert cover letter writer. "
        "Given the following CV and job description, produce exactly this structure:\n\n"
        "COMPANY: <hiring company name in 1-3 words max, NOT the recruiter or agency name — if the company is unnamed use a short descriptor like 'AI Startup'>\n"
        "POSITION: <job title in 3 words max>\n"
        "---\n"
        "<cover letter>\n\n"
        "Rules for the cover letter:\n"
        "- Email body format only (no subject line)\n"
        "- Greeting: use the hiring manager's real name if present; if only a recruiter name is mentioned, address them by name but clarify you are applying for the role they are recruiting for, NOT at their agency; if no name is available, use 'Dear Hiring Manager,'\n"
        "- 2-3 short paragraphs in formal business English\n"
        "- Paragraph 1: state the role and one sharp reason you are a strong fit, referencing something specific about the hiring company (not the recruiter agency)\n"
        "- Paragraph 2: highlight 2-3 concrete skills or experiences from the CV that directly match the key requirements and specific tools listed in the job description\n"
        "- Paragraph 3: one sentence expressing genuine interest and availability; no generic enthusiasm clichés\n"
        "- End with 'Kind regards,\\nLorenzo Schiroli'\n"
        "- No bullet points, no bold text, no placeholders, no copy-pasted phrases from the job description\n\n"
        f"CV:\n{cv_text}\n\n"
        f"Job Description:\n{job_description}"
    )
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
    )
    raw = response.choices[0].message.content.strip()

    company = re.search(r"^COMPANY:\s*(.+)$", raw, re.MULTILINE)
    position = re.search(r"^POSITION:\s*(.+)$", raw, re.MULTILINE)
    company = company.group(1).strip() if company else "company"
    position = position.group(1).strip() if position else "position"

    letter = raw.split("---", 1)[-1].strip()
    return letter, company, position


def generate_linkedin_messages(cv_text: str, job_description: str, client: Groq) -> list[str]:
    prompt = (
        "You are an expert at LinkedIn outreach. Write 3 different LinkedIn connection request messages "
        "from Lorenzo Schiroli to someone who works at the hiring company for the role below.\n\n"
        "Rules for each message:\n"
        "- Maximum 200 characters including spaces\n"
        "- Do not start with 'Hi' or 'Hello' — open with something more direct and specific\n"
        "- Reference the specific role and company naturally\n"
        "- Do not ask for a job directly — ask for a brief chat or to connect\n"
        "- Sound human, not templated\n"
        "- No emojis, no hashtags\n"
        "- Do not mention the cover letter or that you already applied\n"
        "- Each message should have a clearly different tone or angle\n\n"
        "Output format — exactly this, nothing else:\n"
        "1. <message>\n"
        "2. <message>\n"
        "3. <message>\n\n"
        f"CV:\n{cv_text}\n\n"
        f"Job Description:\n{job_description}"
    )
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
    )
    raw = response.choices[0].message.content.strip()
    messages = []
    for match in re.finditer(r"^\d+\.\s*(.+)$", raw, re.MULTILINE):
        messages.append(match.group(1).strip().strip('"')[:200])
    return messages


def _collapse(para, text: str) -> None:
    runs = para.runs
    if not runs:
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def _rm(para) -> None:
    para._element.getparent().remove(para._element)


def generate_cv_modifications(paragraphs: list[tuple[int, str]], job_description: str, client: Groq) -> list[dict]:
    para_list = "\n".join(f"{i}: {text}" for i, text in paragraphs)
    prompt = (
        "You are a CV editor. Your job is to tailor the CV below to the job description.\n\n"
        "You will return a JSON array of edits. Each edit is either:\n"
        '  {"index": N, "action": "replace", "text": "rewritten paragraph"}\n'
        '  {"index": N, "action": "remove"}\n\n'
        "GOLDEN RULE: every word in your replacements must already exist somewhere in the CV. "
        "Do not add any tool name, technology, or skill that is not already written in the CV paragraphs. "
        "If a tool is not in the CV, do not mention it — not even once.\n\n"
        "What you should do:\n"
        "- Rewrite the summary to lead with deployment and MLOps experience, keeping only what is already there\n"
        "- Expand the 1-2 most relevant experience bullets with more detail, using only facts already in the CV\n"
        "- Remove bullets or sections that are clearly irrelevant to this role\n"
        "- Keep all names, dates, company names, contact info, and section headers exactly as they are\n\n"
        "What you must never do:\n"
        "- Invent or add any tool, framework, or technology not already present in the CV\n"
        "- Add prose sentences inside a skills keyword line\n"
        "- Return anything other than the JSON array\n\n"
        f"CV:\n{para_list}\n\n"
        f"Job Description:\n{job_description}"
    )
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
    )
    raw = response.choices[0].message.content.strip()
    raw = re.sub(r"```(?:json)?\s*|\s*```", "", raw).strip()
    return json.loads(raw)


def _cleanup_orphaned_sections(doc) -> None:
    def is_section_header(p) -> bool:
        t = p.text.strip()
        return len(t) > 2 and t == t.upper() and t.replace(" ", "").isalpha()

    paras = list(doc.paragraphs)
    header_indices = [i for i, p in enumerate(paras) if is_section_header(p)]

    to_remove = set()
    for hi, header_i in enumerate(header_indices):
        next_header_i = header_indices[hi + 1] if hi + 1 < len(header_indices) else len(paras)
        content_found = any(
            paras[j].text.strip()
            and not is_section_header(paras[j])
            and not (paras[j].text.strip().endswith(":") and len(paras[j].text.strip()) <= 50)
            for j in range(header_i + 1, next_header_i)
        )
        if not content_found:
            for j in range(header_i, next_header_i):
                if paras[j].text.strip():
                    to_remove.add(id(paras[j]))

    for para in list(doc.paragraphs):
        if id(para) in to_remove:
            _rm(para)


def apply_cv_modifications(cv_path: str, modifications: list[dict], output_path: str) -> None:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copy(cv_path, output_path)
    doc = Document(output_path)
    paragraphs = doc.paragraphs
    for mod in modifications:
        idx = mod.get("index")
        if idx is None or not (0 <= idx < len(paragraphs)):
            continue
        if mod["action"] == "replace":
            _collapse(paragraphs[idx], mod["text"])
        elif mod["action"] == "remove":
            _rm(paragraphs[idx])
    _cleanup_orphaned_sections(doc)
    doc.save(output_path)


def slugify(text: str, max_length: int = 30) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return slug[:max_length].rstrip("-")


def write_cover_letter(text: str, output_path: str) -> None:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()
    for para in text.split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)
    doc.save(output_path)


def main():
    client = get_client()
    cv_text = extract_cv_text(CV_PATH)
    print("CV loaded.")
    job_description = read_job_description()
    print("Job description read from clipboard.")
    print("Generating cover letter...")
    letter, company, position = generate_cover_letter(cv_text, job_description, client)
    folder = f"output/{slugify(company)}-{slugify(position)}"
    output_path = f"{folder}/cover_letter.docx"
    write_cover_letter(letter, output_path)
    print(f"Cover letter saved → {output_path}")
    print("\nGenerating tailored CV...")
    doc = Document(CV_PATH)
    paragraphs = [(i, p.text) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    modifications = generate_cv_modifications(paragraphs, job_description, client)
    cv_output_path = f"{folder}/curriculum_vitae.docx"
    apply_cv_modifications(CV_PATH, modifications, cv_output_path)
    print(f"Tailored CV saved → {cv_output_path}")
    print("\nGenerating LinkedIn messages...")
    linkedin_messages = generate_linkedin_messages(cv_text, job_description, client)
    print("\n--- LinkedIn messages ---")
    for i, msg in enumerate(linkedin_messages, 1):
        print(f"\n{i}. ({len(msg)} chars)\n{msg}")
    print("\n---")


if __name__ == "__main__":
    main()
