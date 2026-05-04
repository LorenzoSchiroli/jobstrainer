"""Generate tailored CV versions: LLM-focused and Computer Vision-focused."""
import shutil
import os
from docx import Document


def collapse(para, text):
    """Put all text in the first run, clear the rest."""
    runs = para.runs
    if not runs:
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def set_skill(para, category, content):
    """Set skill line: run 0 = category name, run 1 = content, rest cleared."""
    runs = para.runs
    if not runs:
        return
    runs[0].text = category
    if len(runs) > 1:
        runs[1].text = content
        for r in runs[2:]:
            r.text = ""
    else:
        runs[0].text = category + content


def rm(para):
    para._element.getparent().remove(para._element)


# ---------------------------------------------------------------------------
# LLM FOCUSED
# ---------------------------------------------------------------------------
src = "data/lorenzo_schiroli_cv.docx"
shutil.copy(src, "data/lorenzo_schiroli_cv_llm.docx")
d = Document("data/lorenzo_schiroli_cv_llm.docx")
p = d.paragraphs

# Title (para 2): runs 0/1/2 are all blue 00B0F0
collapse(p[2], "NLP & LLM Engineer")

# Summary (para 4)
collapse(
    p[4],
    "Machine Learning Engineer with 4 years of experience specializing in NLP and Large Language Models, "
    "with deep expertise in document understanding systems. Experienced across the full ML lifecycle: "
    "fine-tuning, training, and deploying language models in production environments on AWS. "
    "Strong focus on model optimization techniques (LoRA, quantization), practical performance, "
    "and real-world inference cost constraints.",
)

# Skills: swap NLP and CV so NLP comes first (para 15 → NLP, para 16 → CV)
set_skill(
    p[15],
    "NLP",
    ": Transformers, Hugging Face, LLM, BERT, fine-tuning (LoRA, GGUF), RAG, NER, NLTK, Gensim. ",
)
set_skill(
    p[16],
    "Computer Vision",
    ": OpenCV, OCR, VLM, document information extraction, CNN, ViT. ",
)

# Farm4trade company description (para 24)
collapse(
    p[24],
    "Startup building AI document processing and vision-language solutions for industrial applications.",
)

# Farm4trade OCR bullet (para 28): reframe with transformer angle
collapse(
    p[28],
    "Redesigned the OCR pipeline using transformer-based text recognition, raising accuracy from 75% to 95%.",
)

# Farm4trade image segmentation (para 29): remove — not LLM relevant
seg_para = p[29]

# Farm4trade VLM/LoRA (para 30): enhance for LLM audience
collapse(
    p[30],
    "Trained and deployed a Vision-Language Model (VLM) for transport document information extraction, "
    "applying LoRA fine-tuning and GGUF quantization to meet production latency and cost targets.",
)

# Expert.ai: Table Structure Recognition (para 36) → document layout understanding angle
collapse(
    p[36],
    "Developed a document layout analysis model (Table Structure Recognition), improving accuracy from 40% to 80% "
    "through architecture improvements and framework migration.",
)

# Expert.ai: LLM for key-value (para 37): enhance
collapse(
    p[37],
    "Designed and trained a multimodal LLM pipeline for structured key-value extraction from business documents, "
    "combining visual layout features with language understanding.",
)

# Projects section (para 40 heading, 41 intro, 42 Sudoku): remove — Sudoku is CV, not LLM
proj_heading = p[40]
proj_intro = p[41]
sudoku_para = p[42]

# Remove collected paragraphs
rm(seg_para)
rm(proj_heading)
rm(proj_intro)
rm(sudoku_para)

d.save("data/lorenzo_schiroli_cv_llm.docx")
print("✓ LLM version saved → data/lorenzo_schiroli_cv_llm.docx")


# ---------------------------------------------------------------------------
# COMPUTER VISION FOCUSED
# ---------------------------------------------------------------------------
shutil.copy(src, "data/lorenzo_schiroli_cv_cv.docx")
d2 = Document("data/lorenzo_schiroli_cv_cv.docx")
p2 = d2.paragraphs

# Title
collapse(p2[2], "Computer Vision Engineer")

# Summary
collapse(
    p2[4],
    "Machine Learning Engineer with 4 years of experience specializing in Computer Vision and visual "
    "document understanding, working across R&D and production environments. Experienced across the full "
    "ML lifecycle: designing, training, and deploying visual deep learning models on AWS. Strong focus on "
    "image processing pipelines, model performance, and real-world constraints such as latency and "
    "inference cost.",
)

# Skills: enhance CV (para 15), trim NLP (para 16)
set_skill(
    p2[15],
    "Computer Vision",
    ": OpenCV, OCR, VLM, image segmentation, document layout analysis, object detection, CNN, ViT. ",
)
set_skill(
    p2[16],
    "NLP",
    ": Transformers, Hugging Face, BERT, model optimization (LoRA, GGUF), document understanding. ",
)

# Farm4trade company description (para 24): keep — already CV-focused

# Farm4trade OCR (para 28): reinforce CV pipeline angle
collapse(
    p[28],  # Note: using p not p2 — this is intentional (already saved LLM), use p2
    "Redesigned the OCR pipeline with a new computer vision solution, raising accuracy from 75% to 95%.",
)
# Correction — use p2:
collapse(
    p2[28],
    "Redesigned the OCR pipeline with a new computer vision solution, raising accuracy from 75% to 95%.",
)

# Farm4trade image segmentation (para 29): enhance for CV
collapse(
    p2[29],
    "Built a medical image segmentation model for a high-profile client, managing data annotation, "
    "model training, and deployment under strict latency constraints.",
)

# Farm4trade VLM (para 30): emphasize visual understanding
collapse(
    p2[30],
    "Trained and deployed a Vision-Language Model for transport document extraction, optimizing the visual "
    "backbone with LoRA fine-tuning and reducing inference cost via GGUF quantization.",
)

# Expert.ai company description (para 34): acknowledge visual/layout angle
collapse(
    p2[34],
    "Mid-sized NLP company specializing in document intelligence and visual document layout understanding.",
)

# Expert.ai: Table Structure Recognition (para 36): emphasize visual ML
collapse(
    p2[36],
    "Developed a visual deep learning model for Table Structure Recognition, improving accuracy from 40% to 80% "
    "through architecture improvements and framework migration.",
)

# Expert.ai: LLM bullet (para 37): remove — less CV relevant
llm_para = p2[37]

# Expert.ai: reading order (para 38): reframe as document layout analysis
collapse(
    p2[38],
    "Contributed to document layout analysis algorithms for reading order detection in complex documents.",
)

# Sudoku (para 42): already good for CV, leave untouched

rm(llm_para)

d2.save("data/lorenzo_schiroli_cv_cv.docx")
print("✓ CV  version saved → data/lorenzo_schiroli_cv_cv.docx")
