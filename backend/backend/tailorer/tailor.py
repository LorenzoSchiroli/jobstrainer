import io
import json
import logging
import os
import re

import docx
from docx import Document
from groq import AsyncGroq

_LARGE = lambda: os.environ["GROQ_MODEL_LARGE"]
_log = logging.getLogger(__name__)


def _parse_cover_letter_response(raw: str) -> tuple[str, str, str]:
    company = re.search(r"^COMPANY:\s*(.+)$", raw, re.MULTILINE)
    position = re.search(r"^POSITION:\s*(.+)$", raw, re.MULTILINE)
    company = company.group(1).strip() if company else "company"
    position = position.group(1).strip() if position else "position"
    letter = raw.split("---", 1)[-1].strip()
    return company, position, letter


def _build_docx_bytes(text: str) -> bytes:
    doc = Document()
    for para in text.split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _collapse(para, text: str) -> None:
    runs = para.runs
    if not runs:
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def _rm(para) -> None:
    para._element.getparent().remove(para._element)


def _apply_cv_modifications(cv_bytes: bytes, modifications: list[dict]) -> bytes:
    doc = Document(io.BytesIO(cv_bytes))
    paragraphs = doc.paragraphs
    for mod in modifications:
        idx = mod.get("index")
        if idx is None or not (0 <= idx < len(paragraphs)):
            continue
        if mod["action"] == "replace":
            _collapse(paragraphs[idx], mod["text"])
        elif mod["action"] == "remove":
            _rm(paragraphs[idx])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _signoff_rule(applicant_name: str) -> str:
    """Sign with the applicant's profile name, falling back to the CV's own name."""
    if applicant_name:
        return f"- End with 'Kind regards,\\n{applicant_name}'\n"
    return (
        "- End with 'Kind regards,' then a newline and the candidate's full name "
        "exactly as it appears in the CV\n"
    )


async def generate_tailored_documents(
    cv_text: str,
    cv_bytes: bytes,
    job_description: str,
    groq_client: AsyncGroq,
    applicant_name: str = "",
) -> tuple[bytes, bytes, str]:
    """Returns: (tailored_cv_bytes, cover_letter_bytes, cover_letter_text)"""
    cl_prompt = (
        "You are an expert cover letter writer. "
        "Given the following CV and job description, produce exactly this structure:\n\n"
        "COMPANY: <hiring company name in 1-3 words max>\n"
        "POSITION: <job title in 3 words max>\n"
        "---\n"
        "<cover letter>\n\n"
        "Rules:\n"
        "- Email body format only (no subject line)\n"
        "- Greeting: 'Dear Hiring Manager,' if no name is known\n"
        "- 2-3 short paragraphs in formal business English\n"
        f"{_signoff_rule(applicant_name)}"
        "- No bullet points, no bold text, no placeholders\n\n"
        f"CV:\n{cv_text}\n\nJob Description:\n{job_description}"
    )
    cl_resp = await groq_client.chat.completions.create(
        model=_LARGE(),
        messages=[{"role": "user", "content": cl_prompt}],
    )
    raw_cl = cl_resp.choices[0].message.content.strip()
    _, _, cl_text = _parse_cover_letter_response(raw_cl)
    cl_bytes = _build_docx_bytes(cl_text)

    para_list = "\n".join(
        f"{i}: {p.text}"
        for i, p in enumerate(Document(io.BytesIO(cv_bytes)).paragraphs)
        if p.text.strip()
    )
    cv_mod_prompt = (
        "You are a CV editor. Tailor the CV to the job description.\n\n"
        "Return a JSON array of edits. Each edit:\n"
        '  {"index": N, "action": "replace", "text": "rewritten paragraph"}\n'
        '  {"index": N, "action": "remove"}\n\n'
        "GOLDEN RULE: only use words/tools already in the CV. Never invent skills.\n\n"
        f"CV:\n{para_list}\n\nJob Description:\n{job_description}"
    )
    cv_resp = await groq_client.chat.completions.create(
        model=_LARGE(),
        messages=[{"role": "user", "content": cv_mod_prompt}],
    )
    raw_mods = re.sub(r"```(?:json)?\s*|\s*```", "", cv_resp.choices[0].message.content.strip())
    _log.info("[tailor] raw_mods (first 500): %s", raw_mods[:500])
    try:
        modifications = json.loads(raw_mods)
    except json.JSONDecodeError:
        _log.warning("[tailor] failed to parse CV modifications, using original CV")
        modifications = []
    tailored_cv_bytes = _apply_cv_modifications(cv_bytes, modifications)

    return tailored_cv_bytes, cl_bytes, cl_text
